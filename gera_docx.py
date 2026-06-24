#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

AZUL  = RGBColor(0x1a, 0x1a, 0x2e)
AMBER = RGBColor(0xfb, 0xab, 0x48)
BEGE  = RGBColor(0xfa, 0xfa, 0xf7)
CINZA = RGBColor(0x55, 0x55, 0x55)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top='single', bottom='single', left='single', right='single', sz='4', color='DDDDDD'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = AZUL
    # linha laranja abaixo do h1
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), 'fbab48')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg='1a1a2e'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # cabeçalho
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    # linhas
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            if r % 2 == 1:
                set_cell_bg(cell, 'fafaf7')
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    # larguras
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_webinar_block(doc, titulo, data_horario, tema, rows_datas, alerta=None):
    # barra laranja antes do título
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(0)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '24')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), 'fbab48')
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = AZUL

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(4)
    pPr2  = p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    left2 = OxmlElement('w:left')
    left2.set(qn('w:val'),   'single')
    left2.set(qn('w:sz'),    '24')
    left2.set(qn('w:space'), '6')
    left2.set(qn('w:color'), 'fbab48')
    pBdr2.append(left2)
    pPr2.append(pBdr2)
    r2 = p2.add_run(data_horario)
    r2.font.size = Pt(10)
    r2.font.color.rgb = CINZA

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(8)
    r3 = p3.add_run(f'"{tema}"')
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    make_table(
        doc,
        ['Fase', 'D-X', 'Data', 'Semana', 'Formato', 'ADS'],
        rows_datas,
        col_widths=[3.2, 1.4, 1.6, 1.8, 3.5, 1.5]
    )

    if alerta:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_before = Pt(6)
        p4.paragraph_format.space_after  = Pt(4)
        pPr4  = p4._p.get_or_add_pPr()
        pBdr4 = OxmlElement('w:pBdr')
        left4 = OxmlElement('w:left')
        left4.set(qn('w:val'),   'single')
        left4.set(qn('w:sz'),    '12')
        left4.set(qn('w:space'), '6')
        left4.set(qn('w:color'), 'fbab48')
        pBdr4.append(left4)
        pPr4.append(pBdr4)
        r4 = p4.add_run(f'Atenção: {alerta}')
        r4.font.size  = Pt(9)
        r4.font.color.rgb = RGBColor(0x5a, 0x40, 0x00)


doc = Document()

# margens
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# fonte padrão
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── CAPA ───────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run('Webinars Laslo 2026')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = AZUL

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run('Planejamento e Estratégia de Divulgação')
r2.font.size = Pt(14)
r2.font.color.rgb = AMBER

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(20)
r3 = p3.add_run('Laslo Indústria e Comércio Ltda.  ·  Estouro Marketing Digital  ·  Jun/2026')
r3.font.size = Pt(10)
r3.font.color.rgb = CINZA

# ── ESTRATÉGIA ────────────────────────────────────────────────────────────────
heading(doc, 'Estratégia de Divulgação')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Cada webinar recebe uma campanha de ')
r.font.size = Pt(11)
r2 = p.add_run('20 dias')
r2.bold = True
r2.font.size = Pt(11)
r3 = p.add_run(' dividida em duas fases progressivas.')
r3.font.size = Pt(11)

heading(doc, 'Fase 1 — Aquecimento (D-20 a D-10)', level=2)
make_table(doc,
    ['Disparo', 'Formato', 'ADS'],
    [
        ['D-20', 'Story — gancho do tema + "Webinar em 20 dias!"', 'Não'],
        ['D-17', 'Story — data, horário e chamada para inscrição',  'Não'],
        ['D-13', 'Story — pergunta provocativa sobre o tema',        'Não'],
        ['D-10', 'Story — "Faltam 10 dias — link na bio"',          'Não'],
    ],
    col_widths=[2.5, 10.5, 2]
)

heading(doc, 'Fase 2 — Intensificação (D-7 a D-0)', level=2)
make_table(doc,
    ['Disparo', 'Formato', 'ADS'],
    [
        ['D-7', 'Feed — "Salve a data! [data] às 20h — [tema]"',       'SIM'],
        ['D-3', 'Feed — "Inscrições abertas — faltam 3 dias"',          'SIM'],
        ['D-0', 'Feed + Story — "Hoje às 20h — ao vivo! 🔴 link na bio"', 'SIM'],
    ],
    col_widths=[2.5, 10.5, 2]
)

# ── CRONOGRAMA ────────────────────────────────────────────────────────────────
heading(doc, 'Cronograma 2026')
make_table(doc,
    ['Mês', 'Data', 'Horário', 'Produto', 'Tema da Live'],
    [
        ['Junho',    '30/06/2026', '20h', 'ICAVET — Bomba de infusão',       'A "falsa" segurança da infusão: por que sua bomba atual pode estar falhando com pacientes críticos'],
        ['Julho',    '21/07/2026', '20h', 'Descartáveis',                    'Alarme de oclusão: 5 causas que você ignora e que colocam seu paciente em risco'],
        ['Agosto',   '25/08/2026', '20h', 'BS680 — Bomba de seringa',        'TIVA para todos: como uma bomba de seringa democratizou a anestesia de qualidade na veterinária'],
        ['Setembro', '29/09/2026', '20h', 'Elastovet',                       'Seu pós-operatório pode ser domiciliar: controle da dor 24h com infusão contínua, sem gotejamento manual'],
        ['Outubro',  '20/10/2026', '20h', 'BSVET TCI — Alvo controlado',    'Anestesia alvo-controlada: você está pronto para a revolução? O futuro da sua rotina'],
        ['Novembro', '24/11/2026', '20h', 'BVVET',                           'Erros comuns em bombas de infusão: da proteção do paciente a falhas humanas e hidratação de nefropatas'],
        ['Dezembro', '15/12/2026', '20h', 'Pedestal Roama — Ergonomia',      'Centro cirúrgico organizado ou caos de fios: como a organização do espaço influencia o desfecho da anestesia'],
    ],
    col_widths=[2.2, 2.2, 1.8, 4.0, 6.8]
)

# ── DATAS POR WEBINAR ─────────────────────────────────────────────────────────
heading(doc, 'Datas dos Posts por Webinar')

DATAS = [
    {
        'titulo':      'Junho — ICAVET · Bomba de Infusão',
        'data':        '30/06/2026 às 20h',
        'tema':        'A "falsa" segurança da infusão: por que sua bomba atual pode estar falhando com pacientes críticos',
        'rows': [
            ['Aquecimento',    'D-20', '10/06', 'Jun S2', 'Story',          'Não'],
            ['Aquecimento',    'D-17', '13/06', 'Jun S2', 'Story',          'Não'],
            ['Aquecimento',    'D-13', '17/06', 'Jun S3', 'Story',          'Não'],
            ['Aquecimento',    'D-10', '20/06', 'Jun S3', 'Story',          'Não'],
            ['Intensificação', 'D-7',  '23/06', 'Jun S4', 'Feed',           'SIM'],
            ['Intensificação', 'D-3',  '27/06', 'Jun S4', 'Feed',           'SIM'],
            ['Intensificação', 'D-0',  '30/06', 'Jun S5', 'Feed + Story',   'SIM'],
        ],
        'alerta': None,
    },
    {
        'titulo': 'Julho — Descartáveis',
        'data':   '21/07/2026 às 20h',
        'tema':   'Alarme de oclusão: 5 causas que você ignora e que colocam seu paciente em risco',
        'rows': [
            ['Aquecimento',    'D-20', '01/07', 'Jul S1', 'Story',        'Não'],
            ['Aquecimento',    'D-17', '04/07', 'Jul S1', 'Story',        'Não'],
            ['Aquecimento',    'D-13', '08/07', 'Jul S2', 'Story',        'Não'],
            ['Aquecimento',    'D-10', '11/07', 'Jul S2', 'Story',        'Não'],
            ['Intensificação', 'D-7',  '14/07', 'Jul S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-3',  '18/07', 'Jul S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-0',  '21/07', 'Jul S4', 'Feed + Story', 'SIM'],
        ],
        'alerta': 'Julho acumula a antecipação da PET VET Expo (12–14/ago). Priorizar webinar em stories e PET VET em feed.',
    },
    {
        'titulo': 'Agosto — BS680 · Bomba de Seringa',
        'data':   '25/08/2026 às 20h',
        'tema':   'TIVA para todos: como uma bomba de seringa barata democratizou a anestesia de qualidade na veterinária',
        'rows': [
            ['Aquecimento',    'D-20', '05/08', 'Ago S1', 'Story',        'Não'],
            ['Aquecimento',    'D-17', '08/08', 'Ago S2', 'Story',        'Não'],
            ['Aquecimento',    'D-13', '12/08', 'Ago S2', 'Story',        'Não'],
            ['Aquecimento',    'D-10', '15/08', 'Ago S3', 'Story',        'Não'],
            ['Intensificação', 'D-7',  '18/08', 'Ago S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-3',  '22/08', 'Ago S4', 'Feed',         'SIM'],
            ['Intensificação', 'D-0',  '25/08', 'Ago S4', 'Feed + Story', 'SIM'],
        ],
        'alerta': 'D-13 (12/08) coincide com o Dia 1 da PET VET Expo. Publicar o story nos intervalos da feira.',
    },
    {
        'titulo': 'Setembro — Elastovet',
        'data':   '29/09/2026 às 20h',
        'tema':   'Seu pós-operatório pode ser domiciliar: controle da dor 24 horas com infusão contínua, sem gotejamento manual e sem fios',
        'rows': [
            ['Aquecimento',    'D-20', '09/09', 'Set S2', 'Story',        'Não'],
            ['Aquecimento',    'D-17', '12/09', 'Set S2', 'Story',        'Não'],
            ['Aquecimento',    'D-13', '16/09', 'Set S3', 'Story',        'Não'],
            ['Aquecimento',    'D-10', '19/09', 'Set S3', 'Story',        'Não'],
            ['Intensificação', 'D-7',  '22/09', 'Set S4', 'Feed',         'SIM'],
            ['Intensificação', 'D-3',  '26/09', 'Set S4', 'Feed',         'SIM'],
            ['Intensificação', 'D-0',  '29/09', 'Set S5', 'Feed + Story', 'SIM'],
        ],
        'alerta': 'CBAV Salvador acontece em setembro (data a confirmar). Verificar complementaridade com o tema Elastovet.',
    },
    {
        'titulo': 'Outubro — BSVET TCI · Bomba Alvo Controlado',
        'data':   '20/10/2026 às 20h',
        'tema':   'Anestesia alvo-controlada: você está pronto para a revolução? Entenda porque esta anestesia é o futuro da sua rotina',
        'rows': [
            ['Aquecimento',    'D-20', '30/09', 'Set S5', 'Story',        'Não'],
            ['Aquecimento',    'D-17', '03/10', 'Out S1', 'Story',        'Não'],
            ['Aquecimento',    'D-13', '07/10', 'Out S2', 'Story',        'Não'],
            ['Aquecimento',    'D-10', '10/10', 'Out S2', 'Story',        'Não'],
            ['Intensificação', 'D-7',  '13/10', 'Out S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-3',  '17/10', 'Out S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-0',  '20/10', 'Out S4', 'Feed + Story', 'SIM'],
        ],
        'alerta': None,
    },
    {
        'titulo': 'Novembro — BVVET',
        'data':   '24/11/2026 às 20h',
        'tema':   'Erros comuns em bombas de infusão: da proteção do seu paciente de falhas humanas à hidratação de pacientes nefropatas',
        'rows': [
            ['Aquecimento',    'D-20', '04/11', 'Nov S1', 'Story',        'Não'],
            ['Aquecimento',    'D-17', '07/11', 'Nov S1', 'Story',        'Não'],
            ['Aquecimento',    'D-13', '11/11', 'Nov S2', 'Story',        'Não'],
            ['Aquecimento',    'D-10', '14/11', 'Nov S2', 'Story',        'Não'],
            ['Intensificação', 'D-7',  '17/11', 'Nov S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-3',  '21/11', 'Nov S3', 'Feed',         'SIM'],
            ['Intensificação', 'D-0',  '24/11', 'Nov S4', 'Feed + Story', 'SIM'],
        ],
        'alerta': 'D-7 (17/11) coincide com o último dia da Feira PetNor (15–17/nov). Usar o estande para anunciar o webinar presencialmente.',
    },
    {
        'titulo': 'Dezembro — Pedestal Roama · Ergonomia',
        'data':   '15/12/2026 às 20h',
        'tema':   'Centro cirúrgico organizado ou caos de fios: como a organização do espaço físico influencia o desfecho da anestesia',
        'rows': [
            ['Aquecimento',    'D-20', '25/11', 'Nov S4', 'Story',        'Não'],
            ['Aquecimento',    'D-17', '28/11', 'Nov S4', 'Story',        'Não'],
            ['Aquecimento',    'D-13', '02/12', 'Dez S1', 'Story',        'Não'],
            ['Aquecimento',    'D-10', '05/12', 'Dez S1', 'Story',        'Não'],
            ['Intensificação', 'D-7',  '08/12', 'Dez S2', 'Feed',         'SIM'],
            ['Intensificação', 'D-3',  '12/12', 'Dez S2', 'Feed',         'SIM'],
            ['Intensificação', 'D-0',  '15/12', 'Dez S3', 'Feed + Story', 'SIM'],
        ],
        'alerta': None,
    },
]

for wb in DATAS:
    add_webinar_block(doc, wb['titulo'], wb['data'], wb['tema'], wb['rows'], wb.get('alerta'))

# ── CONFLITOS ─────────────────────────────────────────────────────────────────
heading(doc, 'Meses com Evento + Webinar Coincidentes')
make_table(doc,
    ['Mês', 'Webinar', 'Evento no mês', 'Ação sugerida'],
    [
        ['Julho',    'Descartáveis (21/07)',  'Antecipação PET VET Expo',       'Webinar em stories, PET VET em feed'],
        ['Agosto',   'BS680 (25/08)',         'PET VET Expo (12–14/ago)',        'Stories de webinar nos intervalos da feira'],
        ['Setembro', 'Elastovet (29/09)',     'CBAV Salvador (a confirmar)',     'Verificar complementaridade dos temas'],
        ['Novembro', 'BVVET (24/11)',         'PETNOR (15–17/nov)',              'Usar estande PetNor para anunciar o webinar'],
    ],
    col_widths=[2.5, 3.5, 4.5, 6.5]
)

# ── RODAPÉ ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Estouro Marketing Digital  ·  Planejamento Laslo Webinars 2026  ·  Atualizado em junho/2026')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

OUT = '/mnt/c/Users/alexa/laslo-reuniao/webinars-2026.docx'
doc.save(OUT)
print(f'Salvo em {OUT}')
